"""
Install the Google AI Python SDK

$ pip install google-generativeai

See the getting started guide for more information:
https://ai.google.dev/gemini-api/docs/get-started/python
"""

import os
import docx
import google.generativeai as genai
from docx.shared import Pt  # For setting font size
from docx.oxml.ns import qn  # For setting East Asian fonts


from docx import Document
from docx.shared import Inches



from dotenv import load_dotenv
load_dotenv()

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Create the model
# See https://ai.google.dev/api/python/google/generativeai/GenerativeModel
generation_config = {
  "temperature": 1,
  "top_p": 0.95,
  "top_k": 64,
  "max_output_tokens": 200,
  "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
  model_name="gemini-1.5-pro",
  generation_config=generation_config,
  system_instruction="You are a chat bot which is used to generate Projects report of huge paragraphs on given topic, your response should be proper and reliable for storing in a word file in proper format of project report. Use Heading 1 for main sections and Heading 2 for subheadings."

  # safety_settings = Adjust safety settings
  # See https://ai.google.dev/gemini-api/docs/safety-settings
)
chat_session = model.start_chat(
  history=[
  ]
)

response = chat_session.send_message(input())

#print(response.text)


document = Document()

document.add_heading('Document Title', 0)
p = document.add_paragraph(response.text)


for run in p.runs:
    run.font.size = Pt(42)  # Set the font size to 12 points




document.save('demo.docx')

doc = docx.Document("demo.docx")
for paragraph in doc.paragraphs:
  text = paragraph.text.replace("*", "")  # Replace asterisks
  text = text.replace("#", "")           # Replace number signs
  paragraph.text = text

  for run in paragraph.runs:

    run.font.size = Pt(12)  # Ensure font size is 12 points



doc.save("demo.docx")  # Overwrite the file with changes









